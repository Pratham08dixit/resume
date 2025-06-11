import os
import fitz  # PyMuPDF
from docx import Document
import streamlit as st
import json
from dotenv import load_dotenv
import google.generativeai as genai
import time

# Load API Keys from .env
load_dotenv()

# Get API key from environment or Streamlit secrets
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY") or st.secrets.get("GEMINI_API_KEY")

if not GEMINI_API_KEY:
    st.error("⚠️ Missing GEMINI_API_KEY. Please add it to your Streamlit secrets or .env file")
    st.info("For Streamlit Cloud: Add GEMINI_API_KEY to your app secrets")
    st.info("For local development: Add GEMINI_API_KEY to your .env file")
    st.stop()

# Configure Gemini AI
genai.configure(api_key=GEMINI_API_KEY)
model = genai.GenerativeModel('gemini-1.5-flash')

# Resume Extraction Functions

def extract_text_from_pdf(pdf_file):
    doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
    return "\n".join(page.get_text() for page in doc)

def extract_text_from_docx(docx_file):
    document = Document(docx_file)
    return "\n".join(p.text for p in document.paragraphs)

def save_docx(text, output_path):
    doc = Document()
    for line in text.strip().splitlines():
        doc.add_paragraph(line)
    doc.save(output_path)

# AI Analysis Functions
def analyze_resume(resume_text):
    """Analyze resume and return structured feedback"""
    prompt = f"""
    Analyze this resume and return ONLY a valid JSON object with exactly these fields:
    - sections_detected: array of strings (sections found in the resume)
    - missing_sections: array of strings (important sections that are missing)
    - well_written_sections: array of strings (sections that are well-written)
    - quality_score: number between 0-100 (overall resume quality)
    - suggestions: array of strings (specific improvement suggestions)

    Resume:
    {resume_text}

    Return ONLY the JSON object, no other text or explanation.
    """

    try:
        response = model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        return f'{{"error": "Failed to analyze resume: {str(e)}"}}'

def improve_resume(resume_text, feedback):
    """Improve resume based on feedback"""
    prompt = f"""
    Improve this resume based on the feedback provided. Keep the candidate's experience intact
    but improve clarity, grammar, formatting, and overall presentation.

    Original Resume:
    {resume_text}

    Feedback Context:
    {feedback}

    Return the improved resume as clean, well-formatted text.
    """

    try:
        response = model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        return f"Error improving resume: {str(e)}"

def get_job_guidance(improved_resume):
    """Provide job search guidance based on resume"""
    prompt = f"""
    Based on this improved resume, provide job search guidance including:
    1) Recommended job titles to search for
    2) Top companies in the candidate's field
    3) Key job boards and websites to use
    4) Networking tips specific to their industry

    Resume:
    {improved_resume}

    Provide structured, actionable advice in markdown format.
    """

    try:
        response = model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        return f"Error generating job guidance: {str(e)}"

# Streamlit UI
st.title("🚀 Resume Analyzer & Job Finder")
#st.caption("AI-powered resume analysis and job search guidance")

uploaded_file = st.file_uploader("Upload Resume (PDF or DOCX)", type=["pdf", "docx"])
run_button = st.button("🚀 Analyze")

if run_button and uploaded_file:
    st.info("⏳ Processing your resume...")

    # Extract resume text
    if uploaded_file.name.endswith(".pdf"):
        resume_text = extract_text_from_pdf(uploaded_file)
    else:
        resume_text = extract_text_from_docx(uploaded_file)

    if not resume_text.strip():
        st.error("Could not extract text from resume.")
        st.stop()

    # Progress bar for better UX
    progress_bar = st.progress(0)
    status_text = st.empty()

    # Step 1: Analyze Resume
    status_text.text("🔍 Analyzing resume...")
    progress_bar.progress(25)

    feedback_raw = analyze_resume(resume_text)

    # Step 2: Improve Resume
    status_text.text("✨ Improving resume...")
    progress_bar.progress(50)

    improved_resume = improve_resume(resume_text, feedback_raw)

    # Step 3: Generate Job Guidance
    status_text.text("💼 Generating job guidance...")
    progress_bar.progress(75)

    job_guidance = get_job_guidance(improved_resume)

    # Complete
    progress_bar.progress(100)
    status_text.text("✅ Analysis complete!")
    time.sleep(1)
    progress_bar.empty()
    status_text.empty()

    # Display Results

    # Resume Feedback
    st.subheader("📊 Resume Feedback")
    try:
        # Try to parse JSON from the raw output
        raw_output = feedback_raw.strip()

        # Remove any markdown code blocks if present
        if raw_output.startswith('```json'):
            raw_output = raw_output.replace('```json', '').replace('```', '').strip()
        elif raw_output.startswith('```'):
            raw_output = raw_output.replace('```', '').strip()

        feedback_data = json.loads(raw_output)

        # Display key metrics in a more user-friendly way
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Quality Score", f"{feedback_data.get('quality_score', 'N/A')}/100")
        with col2:
            st.metric("Sections Detected", len(feedback_data.get('sections_detected', [])))
        with col3:
            st.metric("Missing Sections", len(feedback_data.get('missing_sections', [])))

        # Detailed feedback
        with st.expander("📋 Detailed Analysis", expanded=True):
            st.json(feedback_data)

    except Exception as e:
        st.warning("⚠️ Could not parse JSON feedback. Showing raw output:")
        st.text_area("Raw Feedback", feedback_raw, height=200)
        st.error(f"Error details: {str(e)}")

    # Improved Resume
    st.subheader("✅ Improved Resume")
    st.text_area("Improved Resume", improved_resume, height=300)

    # Download button
    try:
        save_docx(improved_resume, "improved_resume.docx")
        with open("improved_resume.docx", "rb") as f:
            st.download_button(
                "⬇️ Download Improved Resume",
                f,
                file_name="improved_resume.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    except Exception as e:
        st.error(f"Error creating download file: {str(e)}")

    # Job Search Guidance
    st.subheader("💼 Job Search Guidance")
    st.markdown(job_guidance)

# Sidebar with information
with st.sidebar:
    st.header("📝 Instructions")
    st.write("""
    1. Upload your resume (PDF or DOCX)
    2. Click "Analyze"
    3. Review the feedback and suggestions
    4. Download your improved resume
    5. Follow the job search guidance
    """)

    st.header("🔧 Setup")
    st.write("""
    **For Streamlit Cloud:**
    Add `GEMINI_API_KEY` to your app secrets

    **For Local Development:**
    Create a `.env` file with:
    ```
    GEMINI_API_KEY=your_api_key_here
    ```
    """)
